# -*- coding: utf-8 -*-
"""
Created on Tue May 15 11:47:50 2018

@author: Wei Deng

back up version with no username and pwd
"""


import datetime
import pytz
import atws
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


class clientList:
    def __init__(self):
        self.path="ClientContact.csv"
        self.accountlist=self.cleanlist(self.path)
        self.accounts=self.accountlist['CUSTOMER'].unique()
    
    def cleanlist(self,path ):
        alist=pd.read_csv(path, encoding='cp1252')
        col=['CUSTOMER', 'EMAIL', 'SMSHighPriority', 'SMSName', 'SMSNumber']
        CleanList=pd.DataFrame(columns=col)
        
        for i in range(0,len(alist)):
            CleanList.loc[str(i)]=[str(alist.loc[i]['CUSTOMER']).strip(),str(alist.loc[i]['EMAIL']).strip(), str(alist.loc[i]['SMSHighPriority']).strip(), str(alist.loc[i]['SMSName']).strip(), str(alist.loc[i]['SMSNumber']).strip() ]
            i=i+1
        
        return CleanList
    
    def test(self):
        for a in self.accounts:
            print("accout:", a, "  email:",self.accountlist[self.accountlist['CUSTOMER']==a])


class autoTK:
    
    def __init__(self,Month_Start, Month_End):
       
       
       self.R_START="2018-"+Month_Start+"-01 12:00am"
       #self.R_START="2018-"+Month_Start+"-25 12:00am"
       
       self.R_END="2018-"+Month_End+"-01 12:00am" 


    def getConnect(self, name='',pwd=''):
        conn=atws.connect(username=name,password=pwd)
        
        return conn



#date format: "yyyy-mm-dd"    
    def getEstDateTime(self,timestr):    
        e_t=datetime.datetime.strptime(timestr, '%Y-%B-%d %I:%M%p')
        
        #if month use 'number' , uncomment next line
        #e_t=datetime.datetime.strptime(timestr, '%Y-%m-%d %I:%M%p')
    
        #if time zone is 'Europe/London', then uncomment below line and comment next line
        #e_time=e_t.astimezone(pytz.timezone('Europe/London'))
        
        #if time zone is 'US/Eastern'
        e_time=e_t.astimezone(pytz.timezone('US/Eastern'))

        return e_time
    

    def getTicketListByDate(self,conn,startDate, endDate):
        tk=atws.Query('Ticket')
        tk.WHERE('CreateDate',tk.GreaterThanorEquals,startDate)
        tk.open_bracket('AND')
        tk.WHERE('CreateDate',tk.LessThanOrEquals,endDate)
        tk.close_bracket()
        tklist=conn.query(tk).fetch_all()
        
        return tklist


#get picklist for names of:
#IssueType, SubIssueType,Priority,Status,QueueName by ID numbers
#only for entity: Ticket

    def getPicklist(self,conn,entityName):
        plist=conn.picklist.lookup(entityName)
        return plist
    
    """
    read following names from autotask:
        resource,account, contact
        
    """
    def getAutotaskName(self,conn, entityName):
        r=atws.Query(entityName)
        r.WHERE('id',r.NotEqual,'')
        rlist=conn.query(r).fetch_all()
        return rlist
    
    
    
    def getAccountNameByID(self,accountID, AccountList):
        AccountName="Empty Account Name"
        for a in AccountList:
            if str(a['id'])==str(accountID):
                AccountName=a['AccountName']
        return AccountName
    
    def getContactName(self,tkid, contactlist):
        for a in contactlist:
            if int(a['id'])==int(tkid):
                try:
                    contactName=a['FirstName']+' '+a['LastName']
                except:
                    contactName=a['FirstName']
        return contactName
    
    def getTKList_df(self):
    #get report month( last month) WHOLE ticket list from autotast

    #create new empty dataframe
    #FRDateTime: first response date and time
    #SLVDateTime: solved date and time
    
        col=['TKNumber', 'AccountName', 'Priority','Status','ContactName','Source','Issue','SubIssue','FRDateTime','FRDueDateTime','FR_SLAMet','SLVDateTime','SLVDueDateTime','Actual SLA Met Tickets','QueueID']
        temp=pd.DataFrame(columns=col)
        
        
        startDate=self.getEstDateTime(self.R_START)
        endDate=self.getEstDateTime(self.R_END)
        #startDate=str(reportYear)+'-'+reportMonth+'01'
        #endDate=str(reportYear)+'-'+reportMonth+str(lastDay)
        
        at=self.getConnect()
    
    #use resolvedDateTime<REsolvedDueDateTime check met SLA or not
    #use FirstResponseDateTime<FirstResponseDueDateTime check met SLA OR NOT
    
    #get ticket list for report month
        tklist=self.getTicketListByDate(at,startDate,endDate)
    
    #get picklist for name of issue, subissue, status, priority, ticketsource
        plist=self.getPicklist(at,'Ticket')
        
    #get account list from autotask
        accountlist=self.getAutotaskName(at,'Account')
    
    #get contact list from autotask
        contactlist=self.getAutotaskName(at,'Contact')
        
        i=0
        for tk in tklist:
    #frmet/slvmet=0: sla not met, =1: sla met
            frmet=0
            slvmet=0
            QueueID='NoQueueID'
            
            
            
            issue=plist['IssueType'].reverse_lookup(str(tk['IssueType']))
            try:
                subIssue=plist['SubIssueType'].reverse_lookup(str(tk['SubIssueType']))
            except:
                subIssue="TBC"
            plevel=plist['Priority'].reverse_lookup(str(tk['Priority']))
            status=plist['Status'].reverse_lookup(str(tk['Status']))
            source=plist['Source'].reverse_lookup(str(tk['Source']))
            
            
            rlvdatetime='no record'
            rlvdue='no record'
            if ('QueueID' in tk):
                QueueID=tk['QueueID']
            
            #if ('ResolvedDateTime' in tk and 'ResolvedDueDateTime' in tk):
            if ('ResolvedDueDateTime' in tk):
                rlvdue=tk['ResolvedDueDateTime']
                
                if ('ResolvedDateTime' in tk): 
                    rlvdatetime=tk['ResolvedDateTime']
           
                    if tk['ResolvedDateTime']<tk['ResolvedDueDateTime']:
                        slvmet=1
            
            frdatetime='no record'
            frdue='no record'
            if ('FirstResponseDateTime' in tk and 'FirstResponseDueDateTime' in tk):
                frdatetime=tk['FirstResponseDateTime']
                frdue=tk['FirstResponseDueDateTime']
            
                if tk['FirstResponseDateTime']<tk['FirstResponseDueDateTime']:
                    frmet=1
            
    #get contact name, if no contact, 'empty'
            contactname='empty'
            if ('ContactID' in tk):
                contactname=self.getContactName(tk['ContactID'], contactlist)
    #get accountname from autotask
            accountname='empty'
            if ('AccountID' in tk):
                accountname=self.getAccountNameByID(str(tk['AccountID']),accountlist)
    
    #get ticket number
            #tknum=self.getTicketNumberByID(at,tk['id'])
            tknum=tk['TicketNumber']
    
            temp.loc[str(i)]=[str(tknum),str(accountname),str(plevel), str(status),str(contactname),str(source),str(issue),str(subIssue),str(frdatetime),str(frdue),str(frmet),str(rlvdatetime),str(rlvdue), int(slvmet), str(QueueID)]
            i=i+1
        
        #only return service desk queue , queueid=29683528
        tklist_df=temp[(temp['QueueID']=='29683528')]
        
        #remove column: queueID
        tklist_df.drop('QueueID', axis=1, inplace=True)
        
        return tklist_df
            
class CreateDoc:
    def __init__(self):
        self.DayOfRelease='15'
        self.VERSION='1.0'
        self.STATUS='FINAL'
        self.DELIVER="Sean Buckingham"
        self.DEV_MANAGER="Sean Buckingham"
        
        
        self.currentMonth=datetime.datetime.now().month
        self.currentYear=datetime.datetime.now().year
        self.reportYear=self.currentYear
        #self.tklist=tklist_df
        self.title=self.createDocTilte(self.currentMonth,self.currentYear)
        self.releaseDate=self.createReleaseDate()
        self.reportMonth,self.lastDay=self.getMonth(self.currentMonth)
        
        self.reportStart="01/"+self.reportMonth+"/"+str(self.reportYear)
        self.reportEnd=str(self.lastDay)+"/"+self.reportMonth+"/"+str(self.reportYear)
    
    def getMonth(self,m):
        if m==1:
            r='December'
            e=31
        if m==2:
            r='January'
            e=31
        if m==3: 
            r='February'
            e=28
        if m==4: 
            r='March'
            e=31
        if m==5: 
            r='April'
            e=30
        if m==6: 
            r='May'
            e=31
        if m==7: 
            r='June'
            e=30
        if m==8: 
            r='July'
            e=31
        if m==9: 
            r='August'
            e=31
        if m==10: 
            r='Septmeber'
            e=31
        if m==11: 
            r='October'
            e=31
        if m==12: 
            r='November'
            e=30
        
        reportMonth=r
        lastdayofMonth=e
        
        return  reportMonth, lastdayofMonth    
    
    
    
    def countMonthTicket(self,accountname,tklist_df):
    
        total=tklist_df[(tklist_df['AccountName']==accountname)]
        completed=total[(total['Status']=='Complete')]
        #completed=total[(total['SLVDateTime']!='no record')]
        
        completedNum=len(completed)
        createNum=len(total)
        
        return completedNum,createNum

    def createDocTilte(self,monthNumber,reportYear):
        reportMonth,lastDay=self.getMonth(monthNumber)
        title="Delivery of Service - "+str(reportMonth)+","+str(reportYear)
        
        return title
    
    def createReleaseDate(self):
        
        release=self.DayOfRelease+"/"+str(self.currentMonth)+"/"+str(self.currentYear)
        
        return release
    
    
    def IssueGraph(self,accountname, tklist):
    #draw bar chart by tickek count amount vs. ticket issue types
        
        tk=tklist[(tklist['AccountName']==accountname)]
        tickets=tk[(tk['Issue']!='Service Desk')]
        issuecount=tickets.groupby(['Issue']).agg(['count'])
        #issuecount.sort_index(inplace=True)
        topissue=issuecount['Status'].sort_values(by=['count'], ascending=False).index[0]
        
        #issue name with max ticket counts
        max_count=issuecount['Status'].sort_values(by=['count'], ascending=False).max()
        
        ticks=self.getMaxTick(int(max_count))
        
        
        img=issuecount['Status'].sort_values(by=['count'], ascending=False).head().plot(kind='barh', xticks=ticks,title="Top Five Issue Types", legend=False)
        img.set_xlabel("Count of Tickets")
        
        #add data lable to bars
        for p in img.patches:
            img.annotate(str(int(p.get_width())),(p.get_x()+p.get_width(),p.get_y()))
            
        return topissue,img
    
    def SubIssueGraph(self,accountname, tklist):
        #draw bar chart by tickek count amount vs. ticket subissue types
        
        tk=tklist[(tklist['AccountName']==accountname)]
        tickets=tk[(tk['SubIssue']!='New Ticket')]
        issuecount=tickets.groupby(['SubIssue']).agg(['count'])
        #issuecount.sort_index(inplace=True)
        topsubissue=issuecount['Status'].sort_values(by=['count'], ascending=False).index[0]
        
        #name of sub issue with max ticket counts
        max_count=issuecount['Status'].sort_values(by=['count'], ascending=False).max()
        
        ticks=self.getMaxTick(int(max_count))
        
        img=issuecount['Status'].sort_values(by=['count'], ascending=False).head().plot(kind='barh', xticks=ticks, title="Top Five SubIssue Types", legend=False)
       
        
        
        img.set_xlabel("Count of Tickets")
        
        #add data lable to bars
        for p in img.patches:
            img.annotate(str(int(p.get_width())),(p.get_x()+p.get_width(),p.get_y()))
        
        return topsubissue,img
    
    def TopUsersGraph(self,accountname,tklist):
        tk=tklist[(tklist['AccountName']==accountname)]
        tickets=tk[(tk['ContactName']!='empty')]
        issuecount=tickets.groupby(['ContactName']).agg(['count'])
        #issuecount.sort_index(inplace=True)
        topuser=issuecount['Status'].sort_values(by=['count'], ascending=False).index[0]
        
        max_count=issuecount['Status'].sort_values(by=['count'], ascending=False).max()
        
        ticks=self.getMaxTick(int(max_count))
        
        img=issuecount['Status'].sort_values(by=['count'], ascending=False).head().plot(kind='barh', title="Top Five Users ", xticks=ticks,legend=False)
        img.set_xlabel("Count of Tickets")
        
        
        for p in img.patches:
            img.annotate(str(int(p.get_width())),(p.get_x()+p.get_width(),p.get_y()))
        
        #img.xaixs.set_major_locator(MaxNLocator(integer=True))
        
        return topuser,img
    
    def TicketPriorityGraph(self,accountname,tklist):
        tickets=tklist[(tklist['AccountName']==accountname)]
        tkcount=tickets.groupby(['Priority']).agg(['count'])
        #issuecount.sort_index(inplace=True)
        #topissue=issuecount['Status'].sort_values(by=['count'], ascending=False).index[0]
        img=tkcount['Status'].sort_values(by=['count'], ascending=False).plot(kind='bar', rot=3,title="Ticket Counts by Priority", legend=False)
        img.set_ylabel("Count of Tickets")
        
        
        return img
    
    
    def getMaxTick(self,count):
        #calculate max tick depending on count number
        for i in range(1000,10,-10):
            if (i>count):
                temp=i
        count=temp
        ticks=range(0,int(count)+5,10)
        
        return ticks
    
    
    def sla_Cal(self,client,tklist):
    
    #change sla-met columns to int
    #tbs.groupby(['Issue','SubIssue','SLAMet'],as_index=False).sum()
    #tbs.groupby(by=['Issue','SubIssue'])['Actual SLA Met Tickets'].count()
    
        #col=['Issue','SubIssue','Total Tickets','SLA Met Tickets','Actual % of SLA Met']
        #slaByIssue=pd.DataFrame(columns=col)
        
        total=tklist[(tklist['AccountName']==client)]
        clientlist=total[(total['SLVDateTime']!='no record')]
        
        
        #.sum: get totalmet sla ticket numbers, becuase if not met, value=0
        #.count: get total ticket numbers,no matter met sla (value=1) or not met(value=1)
        slaByIssue=clientlist.groupby(by=['Issue','SubIssue'], as_index=False).sum()
        temp=clientlist.groupby(by=['Issue','SubIssue'], as_index=False).count()
        slaByIssue['Total Tickets']=temp['TKNumber']
        try:
            slaByIssue['Actual % of SLA Met']=slaByIssue['Actual SLA Met Tickets']/slaByIssue['Total Tickets']
        except:
            slaByIssue['Actual % of SLA Met']=0
        
        slaByIssue['Actual % of SLA Met']=slaByIssue['Actual % of SLA Met'].map('{:.1%}'.format)
        slaByIssue['Actual SLA Met Tickets']=slaByIssue['Actual SLA Met Tickets'].map('{0:4.0f}'.format)
        #col=['Priority', 'SLA Met Tickets', 'Total Tickets','Actual % of SLA Met' ]
        #slaByPriority=pd.DataFrame(columns=col)
        
        slaByPriority=clientlist.groupby(by=['Priority'], as_index=False).sum()
        temp=clientlist.groupby(by=['Priority'], as_index=False).count()
        slaByPriority['Total Tickets']=temp['TKNumber']
        try:
            slaByPriority['Actual % of SLA Met']=slaByPriority['Actual SLA Met Tickets']/slaByPriority['Total Tickets']
        except:
            slaByPriority['Actual % of SLA Met']=0
        
        slaByPriority['Actual % of SLA Met']=slaByPriority['Actual % of SLA Met'].map('{:.1%}'.format)
        slaByPriority['Actual SLA Met Tickets']=slaByPriority['Actual SLA Met Tickets'].map('{0:4.0f}'.format)
        
        
        tb_slaByIssue=slaByIssue[['Issue','SubIssue','Total Tickets','Actual SLA Met Tickets','Actual % of SLA Met']]
        return tb_slaByIssue, slaByPriority           
    
    def sla_fail(self,client, tklist):
        client_completed_tk=tklist[(tklist['Status']=='Complete')]
        SLA_Fail_01=client_completed_tk[(client_completed_tk['AccountName']==client)]
        SLA_Fail_Full=SLA_Fail_01[(SLA_Fail_01['Actual SLA Met Tickets']==0)]
        SLA_Fail=SLA_Fail_Full[['TKNumber','Issue','SubIssue']]
        SLA_Fail['Description']='-'
        return SLA_Fail
        
    def writeDoc(self,Client, table_issuesla, table_plevelsla,tklist,contactlist):
    
        doc=Document()
        
        style=doc.styles['Normal']
        font=style.font
        font.name='Calibri'
        font.size=Pt(12)
        
        doc.add_picture("report_coverImg.png", width=Cm(16), height=Cm(24))
        
        doc.add_heading('GENERAL REPORT OVERVIEW',1)
        para_01="The Delivery of Service document is a monthly  report prepared by Tickbox presenting summary data for "+Client+ ". This activity is from the period of "+self.reportStart+" -- "+self.reportEnd+"."
        p=doc.add_paragraph(para_01)
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    
        p=doc.add_paragraph('For questions or feedback regarding this monthly  report, please contact sean.buckingham@tickbox.com.au')
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    
        
        p=doc.add_paragraph('')
        doc.add_heading('DOCUMENT CONTROL',1)
        table=doc.add_table(rows=6, cols=2)
        #table.style='Table Grid'
        table.style='Light Shading Accent 1'
        table.cell(0,0).text="File Name:"
#        title=createDocTilte()
        table.cell(0,1).text=self.title
        
        table.cell(1,0).text="Version:"
        table.cell(1,1).text=self.VERSION
        #table.cell(1,1).text="V1.0"
        
        table.cell(2,0).text="Status:"
        table.cell(2,1).text=self.STATUS
        #table.cell(2,1).text= "FINAL"
        
        table.cell(3,0).text="Release Date:"
        #releaseDate=createReleaseDate()
        table.cell(3,1).text=self.releaseDate
        
        table.cell(4,0).text="Delive From:"
        table.cell(4,1).text=self.DELIVER
        #table.cell(4,1).text="Sean Buckingham"
        table.cell(5,0).text="Report Manager:"
        table.cell(5,1).text=self.DEV_MANAGER
        #table.cell(5,1).text="Sean Buckingham"
        
        p=doc.add_paragraph('')
        doc.add_heading('EMERGENCY SMS KEY CONTACT LIST',1)
        
        #contactlist=pd.read_csv("ClientContact.csv", encoding='cp1252')
        shortlist=contactlist[(contactlist['CUSTOMER']==Client)]
        row=len(shortlist)
        smstable=doc.add_table(rows=row+1, cols=3)
    
        #smstable.style='Table Grid'
        smstable.style='Light Shading Accent 1'
        
        smstable.cell(0,0).text="Contact Name"
        smstable.cell(0,1).text="Telephone"
        smstable.cell(0,2).text="Email"
        
        pos=shortlist.index
        
        for r in range(row):
            try:
                smstable.cell(r+1,0).text=shortlist['SMSName'][pos[r]]
            except:
                smstable.cell(r+1,0).text='TBC'
            try:
                smstable.cell(r+1,1).text=shortlist['SMSNumber'][pos[r]]
            except:
                smstable.cell(r+1,1).text='TBC'
            smstable.cell(r+1,2).text=shortlist['EMAIL'][pos[r]]
        
        doc.add_page_break()
    #    doc.add_heading('Executive Summary',1)
    #    doc.add_paragraph()
        
        doc.add_heading('Tickbox Service Desk',1)
        doc.add_heading('Service Desk Overview',2)
        p=doc.add_paragraph('The Tickbox Service Desk is committed to ensuring your requests are responded to and resolved within our SLA targets.')
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    
        tk_complete,tk_count=self.countMonthTicket(Client,tklist)
        
        p=doc.add_paragraph('From '+self.reportStart+' to '+self.reportEnd+', Tickbox received total of '+str(tk_count)+' tickets from '+Client+', and completed '+str(tk_complete)+' tickets.')
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    
    
        doc.add_heading('High Priority Ticket',2)
        ptext="Incidents cause interruptions of business activities and must be solved with certain urgency. The following addresses high priority 1 and priority 2 incident tickets, their cause, resolution, feedback, date of occurrence and duration time."
        p=doc.add_paragraph(ptext)
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        
    #    doc.add_heading("Total Ticket Amounts by Priorities",4)
    #    y=TicketPriorityGraph(Client,tklist)
    #    x=y.get_figure()
    #    filename=Client+".png"
    #    x.savefig(filename,bbox_inches='tight')
    #    doc.add_picture(filename,height=Cm(8))
        doc.add_page_break()
        
        doc.add_heading('High Priority Ticket Response and Resolution',2)
        table=doc.add_table(rows=3, cols=4)
        table.style='Light Shading Accent 1'
        table.cell(0,0).text="Priority Level"
        table.cell(0,1).text="Interpretation"
        table.cell(0,2).text="Response Target"
        table.cell(0,3).text="Resolution Target"
        table.cell(1,0).text="Priority 1"
        table.cell(2,0).text="Priority 2"
        
        table.cell(1,1).text="High impact to customers identified as severe incident resulting in an outage."
        table.cell(1,2).text="Response Target within 15 minutes"
        table.cell(1,3).text="Resolution Target within 4 hours"
        
        table.cell(2,1).text="High business impact to customers identified as potential loss or interruption of service."
        table.cell(2,2).text="Response Target within 30 minutes"
        table.cell(2,3).text="Resolution Target within 8 hours"
        
        doc.add_heading('Classification of New Tickets',2)
        ptext="A Tickbox Engineer categorises each ticket based on an Issue and Sub-Issue type to assist the Service Desk in solving the issue and to report back to our clients on those numbers."
        p=doc.add_paragraph(ptext)
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_heading("Issue Type by Ticket Numbers",3)
        
        issnum,yissue=self.IssueGraph(Client,tklist)
        subnum,ysub=self.SubIssueGraph(Client,tklist)
        ptext="The top issue type this month was \""+str(issnum)+"\" request and the top sub issue type was \""+subnum+"\" request."
        p=doc.add_paragraph(ptext)
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        
        
        doc.add_heading("Top Five Issues by Amount of Tickets",4)
        
        x=yissue.get_figure()
        filename=Client+".png"
        x.savefig(filename,bbox_inches='tight')
        doc.add_picture(filename,height=Cm(8))
        doc.add_page_break()
        
        doc.add_heading("Top Five SubIssues by Amount of Ticketss",4)
        
        x=ysub.get_figure()
        filename=Client+".png"
        x.savefig(filename,bbox_inches='tight')
        doc.add_picture(filename,height=Cm(8))
        doc.add_page_break()
        
        doc.add_heading("Additional Ticket Information",3)
        usernum,y=self.TopUsersGraph(Client,tklist)
    
        ptext="This month, total "+usernum+" different users logged cases."
        
        doc.add_heading("Most Active Users",4)
        x=y.get_figure()
        filename=Client+".png"
        x.savefig(filename,bbox_inches='tight')
        doc.add_picture(filename,height=Cm(8))
        doc.add_page_break()
        
    
        #table_issuesla,table_plevelsla
        
        
        title="SLA Met Result"
        doc.add_heading(title,3)
        
        sla_mettotal=0
        for i in range(table_plevelsla.shape[0]):
            sla_mettotal=sla_mettotal+int(table_plevelsla['Actual SLA Met Tickets'][i])
            
        
        #d=int(table_plevelsla['Total Tickets'].sum())/int(table_plevelsla.shape[0])
        #sla_actual=float(table_plevelsla['Actual SLA Met Tickets'].sum())/d
        sla_total=int(table_plevelsla['Total Tickets'].sum())
        sla_actual=sla_mettotal/sla_total
        sla_p="{:.1f}%".format(sla_actual*100)
        #sla_mettotal=int(table_plevelsla['Actual SLA Met Tickets'].sum())
        
        ptext="This month,"+str(int(sla_mettotal))+" out of "+str(int(sla_total))+" tickets met SLA standard, and the SLA met percentage is: "+str(sla_p)+"."
        p=doc.add_paragraph(ptext)
        p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    #    doc.add_page_break()
    
    
    #    tb=doc.add_table(table_plevelsla.shape[0]+1,table_plevelsla.shape[1])
    #    tb.style='Table Grid'
    #    for j in range(table_plevelsla.shape[-1]):
    #        tb.cell(0,j).text=table_plevelsla.columns[j]
    #    
    #    for i in range(table_plevelsla.shape[0]):
    #        for j in range(table_plevelsla.shape[-1]):
    #            tb.cell(i+1,j).text=str(table_plevelsla.values[i,j])
        #doc.add_page_break()
      
        title="Table: SLA Met Result by Issue and SubIssue Types"
        doc.add_heading(title,3)
        tb=doc.add_table(table_issuesla.shape[0]+1,table_issuesla.shape[1])
        tb.style='Table Grid'
        for j in range(table_issuesla.shape[-1]):
            tb.cell(0,j).text=table_issuesla.columns[j]
        
        for i in range(table_issuesla.shape[0]):
            for j in range(table_issuesla.shape[-1]):
                tb.cell(i+1,j).text=str(table_issuesla.values[i,j])
        
        #added at 19-04-2018 for if fail sla, add table for sla fail tickets
        Count_SLA_Fail=self.sla_fail(Client,tklist)
        if (len(Count_SLA_Fail)>0):
            doc.add_page_break()
            title="Table: SLA Failed Ticket List"
            doc.add_heading(title,3)
            tb=doc.add_table(Count_SLA_Fail.shape[0]+1,Count_SLA_Fail.shape[1])
            tb.style='Table Grid'
            
            
            for j in range(Count_SLA_Fail.shape[-1]):
                tb.cell(0,j).text=Count_SLA_Fail.columns[j]
        
            for i in range(Count_SLA_Fail.shape[0]):
                for j in range(Count_SLA_Fail.shape[-1]):
                    tb.cell(i+1,j).text=str(Count_SLA_Fail.values[i,j])
        
        doc.add_page_break()
        doc.add_picture("report_endImg.jpg", height=Cm(24))
        file="./new/"+Client+".docx"
        doc.save(file)


    def GenerateDoc(self,client, tklist,contactlist):
    
        #calculate SLA and return dataframs
        #one for sla by issue, one for sla by priority level
        table_issuesla,table_plevelsla=self.sla_Cal(client,tklist)
        
        #function to create actual word document
        self.writeDoc(client,table_issuesla,table_plevelsla,tklist,contactlist)            
        

def main():
#this is the main function to create monthly report.
  
#intro: when create april report, input april AND may
#for May report, should input May and June
#for Janurary report, CHECK REPORTYEAR NUMBER IN _INIT__():
    
    t=autoTK("June","July")
    
    #get ticket in dataframe format
    tk=t.getTKList_df()
    
    #get client list of monthly report
    clients=clientList()
    
    #class for generate word document
    docs=CreateDoc()
    for c in clients.accounts:
       if len(tk[(tk['AccountName']==c.strip())])>0:
            client=c.strip()
            
            try:
                docs.GenerateDoc(client,tk, clients.accountlist)
                print(c,"report made",)
            except:
                print("please check data of:",c,len(tk[(tk['AccountName']==c.strip())]))
    
    return tk
    
           